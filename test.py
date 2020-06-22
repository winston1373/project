# -*- coding: utf-8 -*-
"""

"""
from nltk.tag import pos_tag
from nltk.tokenize import word_tokenize
from nltk.corpus import wordnet as wn
import docx
import re
from tabula import read_pdf
import PyPDF2 
import time
import random
import pandas as pd
import csv
options = ['yes','no']
string_list = []
options_list = []

def max_key(dic):
    val = list(dic.values())
    keys = list(dic.keys())  

    return keys[val.index(max(val))]


def options_type(options):

	options_counts= {}
	for word in options:
		word_counts = {}
		for synset in wn.synsets(word):
			word_type = synset.lexname().split('.')[0]
			if word_type not in word_counts:
				word_counts[word_type] = 1
			else:
				word_counts[word_type] +=1

		
		print(word, word_counts)
		if len(word_counts) == 0:
			word_type = 'undefined'
		else:
			word_type = max_key(word_counts)

		if word_type not in options_counts:
				options_counts[word_type] = 1
		else:
			options_counts[word_type] +=1

	return max_key(options_counts).upper()

def read_wordfile(filename):
	blank = 'XX'
	doc = docx.Document(filename)
	string_list = []
	options_list = []
	for i in range(len(doc.paragraphs)):
		if i % 2 == 0:
			s = doc.paragraphs[i].text			
			s = re.sub(r'_+', blank, s)
			string_list.append(s)
			# print(s)
		else:

			options = doc.paragraphs[i].text.split('\t')
			print(i, options)
			for j in range(4):
				options[j] = options[j].split()[1]
			options_list.append(options)
			
	return string_list, options_list
	
filename = 'word/102a.docx'

# doc = docx.Document(filename)
# i = 0
# for para in doc.paragraphs:
# 	print(i, para.text)
# 	print()
# 	i += 1
# print(read_wordfile(filename))
# for i in range(101,108+1):
# 	filename = "word/" + str(i) + ".docx"
# 	print(read_wordfile(filename))

# options = ['chronically', 'hysterically', 'simultaneously', 'resistantly']
# options_type(options)


def read_answers(csvname):	
	f = open(csvname, 'r')
	ans = f.read().split(',')
	f.close()
	return ans
csvname = 'answer/101.csv'
answer = read_answers(csvname)



def build_result_dict(results):
	op = ['A', 'B', 'C', 'D']
	result_dict = {}
	for i in range(len(results)):
		result_dict[op[i]] = results[i]
	return result_dict

acc = 0
all_ans_acc = 0
correct_ans_acc = 0
f = open('human.txt', 'r')
lines = f.readlines()
f.close()
n_problems = len(lines)
for line in lines:
	s = line.split(",")
	results = [int(s[2]), int(s[3]), int(s[4]), int(s[5].rstrip())]
	result_dict = build_result_dict(results)
	ans_pred = max_key(result_dict)
	total = sum(result_dict.values())
	acc_pr = 0
	if total != 0 :
		acc_pr = result_dict[s[1]] / total
	all_ans_acc += acc_pr
	if s[0] == s[1]:
		acc += 1
		correct_ans_acc += acc_pr

	# print(s)
print("accuracy: %f  %d/%d" % (acc / n_problems, acc, n_problems))
print("all_ans_acc: %f" % (all_ans_acc / n_problems))
print("correct_ans_acc: %f" % (correct_ans_acc / acc))