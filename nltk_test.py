# -*- coding: utf-8 -*-
from nltk.tag import pos_tag
from nltk.tokenize import word_tokenize
from nltk.corpus import wordnet as wn
import docx
import re
import requests       #引入requests模組
from lxml import etree
import time
import random
def max_key(dic):
    val = list(dic.values())
    keys = list(dic.keys())
    return keys[val.index(max(val))]


def options_type(options):
	options_counts= {}
	for word in options:
		word_counts = {}
		for synset in wn.synsets(word):
			word_type = synset.lexname().split('.')[0]
			if word_type not in word_counts:
				word_counts[word_type] = 1
			else:
				word_counts[word_type] +=1

		
		# print(word, word_counts)
		if len(word_counts) == 0:
			word_type = 'undefined'
		else:
			word_type = max_key(word_counts)

		if word_type not in options_counts:
				options_counts[word_type] = 1
		else:
			options_counts[word_type] +=1

	return max_key(options_counts).upper()


def find_index(forward_index, backward_index, find_type, back=True, forward=True):
	backward_move = back
	forward_move = forward
	stop_index = -1
	type_index = -1
	find_type_list = []
	find_type_list.append(find_type)
	if find_type == 'NOUN':
		find_type_list.append('PRON')
	while(forward_index >= 0 or backward_index < len(a)):

			if (forward_index >= 0 and 
				forward_move == True
				):
				if(a[forward_index][0] in MARKS):
					stop_index = forward_index + 1
					forward_move = False
				if a[forward_index][1] in find_type_list:
					# print("f",a[forward_index])
					type_index = forward_index
					break;

			if (backward_index < len(a)
				and backward_move == True
				):
				if(a[backward_index][0] in MARKS):
					stop_index = backward_index - 1
					backward_move = False
				if (a[backward_index][1] in find_type_list):
					# print("b",a[backward_index])
					type_index = backward_index
					break
			forward_index = forward_index - 1 
			backward_index = backward_index + 1

	if type_index == -1:
		return stop_index
	return type_index

def substring(first, last):
	substring = ""
	i = first
	while i <= last:
		if i == last:
			substring += a[i][0]
			break
		if (i+2 < len(a) and a[i+2][0] == 's'):
			substring += a[i][0] + a[i+1][0] + 's ' +  a[i+3][0]
			i += 4
			continue

		substring += a[i][0] + " "
		i += 1
	return substring


def read_wordfile(filename):
	blank = 'XX'
	doc = docx.Document(filename)
	string_list = []
	options_list = []
	for i in range(len(doc.paragraphs)):
		if i % 2 == 0:
			s = doc.paragraphs[i].text
			s = re.sub(r'_+', blank, s)
			string_list.append(s)
			# print(s)
		else:
			options = doc.paragraphs[i].text.split('\t')
			for j in range(4):
				options[j] = options[j].split()[1]
			options_list.append(options)
			# print(options)
	return string_list, options_list


def google_search(string, choice):
    
    key = string.split()
    # key=['dangerous','object']

    key_added='+'.join(key)
    #print(key_added)
    url = f"https://www.google.com/search?hl=en&q={key_added}"

    # print("input your choice:")      #choice=0or1  0為一般搜尋 1為進階搜尋
    # choice = int(input())
    #print("%d" %choice )


    #用GET下載網頁
    #一般搜尋
    if choice == 0:
        url = f"https://www.google.com/search?hl=en&q={key_added}"
        #print(url)
    #進階搜尋
    elif choice == 1:
        url = f"https://www.google.com/search?hl=en&q=\"{key_added}\"&oq=\"{key_added}\""
        #print(url)
    #錯誤輸入
    else:
        print("Wrong input")

    #if同時有一般也有進階 eat "apple pie"


    #加上user agent讓網頁識別身分
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/81.0.4044.138 Safari/537.36',
    }

    response = requests.get(url, headers=headers)

    html = response.text
    page = etree.HTML(html)
    num = 0
    if html.find("did not match any documents") > 0 or html.find("No results found for") > 0:
        # print("0 result")
        num = 0

    else:
        #print(page.xpath(u'//*[@id="result-stats"]/text()'))
        
        #print(response.status_code)       #伺服器回應的狀態碼

        result=page.xpath(u'//*[@id="result-stats"]/text()')
        result="".join(result)     #List to string
        result=result.split(" ")
        for i in range(len(result)):
            if result[i].find("result")!=-1:
                num = int(float(result[i-1].replace(',','')))
                # print(result[i-1])
        # print(result)
        # result = int(float(result[1].replace(',','')))

    return num


def ADJ_target(forward_index, backward_index):
	search_string = ""
	VERB_index = find_index(forward_index, backward_index, 'VERB')
	if VERB_index < target_index:
		if a[VERB_index][0] in BeVERB:
			NOUN_index = find_index(forward_index, backward_index, 'NOUN', back=False)
			search_string = substring(NOUN_index, target_index)
		else:
			NOUN_index = find_index(forward_index, backward_index, 'NOUN', forward=False)
			search_string = substring(VERB_index, NOUN_index)
			

	if search_string == "":
		VERB_index = find_index(forward_index, backward_index, 'VERB', forward=False)
		search_string = substring(target_index, VERB_index)


	return search_string


def NOUN_target(forward_index, backward_index):
	search_string = ""
	VERB_index = find_index(forward_index, backward_index, 'VERB')
	if VERB_index < target_index :
		search_string = substring(VERB_index, target_index)
	else:
		ADJ_index = find_index(target_index, backward_index, 'ADJ', back=False)
		# print(ADJ_index)
		if(a[ADJ_index][1] == 'ADJ'):
			search_string = substring(ADJ_index, VERB_index)
		else:
			search_string = substring(target_index, VERB_index)
	return search_string


def VERB_target(forward_index, backward_index):
	search_string = ""
	NOUN_index = find_index(forward_index, backward_index, 'NOUN', forward=False)
	
	if a[NOUN_index][1] != 'NOUN':
		NOUN_index = find_index(forward_index, backward_index, 'NOUN', back=False)
		
	NOUN_index = compound_NOUN(NOUN_index)
	
	if NOUN_index < target_index :
		search_string = substring(NOUN_index, target_index)
	else:
		search_string = substring(target_index, NOUN_index)
	return search_string


def ADV_target(forward_index, backward_index):
	search_string = ""
	step = 1
	while True:
		if (a[target_index+step][1] == 'VERB' and a[target_index+step][1] not in BeVERB):
			backward_index += step
			search_string = VERB_target(forward_index, backward_index)
			break

			
		elif (a[target_index-step][1] == 'VERB'and a[target_index-step][0] not in BeVERB):
			forward_index -= step

			
			search_string = VERB_target(forward_index, backward_index)
			break

		elif (a[target_index+step][1] == 'NOUN'):
			backward_index += step
			print(a[target_index-step][1])
			search_string = NOUN_target(forward_index, backward_index)
			break
			
		elif (a[target_index-step][1] == 'NOUN'):
			forward_index -= step
			search_string = NOUN_target(forward_index, backward_index)
			break
		step += 1
	return search_string


def compound_NOUN(NOUN_index):
	if NOUN_index < target_index :
		while(a[NOUN_index][1] == 'NOUN'):
			if(NOUN_index-1 >= 2 and a[NOUN_index-1][1] == 'NOUN'):
				NOUN_index -= 1
			else:
				break
	else:
		while(a[NOUN_index][1] == 'NOUN'):
			if(NOUN_index+1 < len(a) and a[NOUN_index+1][1] == 'NOUN'):
				NOUN_index += 1
			else:
				break
	return NOUN_index


def read_answers(csvname):	
	f = open(csvname, 'r')
	ans = f.read().split(',')
	f.close()
	return ans


def build_result_dict(results):
	op = ['A', 'B', 'C', 'D']
	result_dict = {}
	for i in range(len(results)):
		result_dict[op[i]] = results[i]
	return result_dict

file = "101"
filename = 'word/{}.docx'.format(file)
ansname  = "answer/{}.csv".format(file)
resultname = "result/{}_result.csv".format(file)

string_list, options_list = read_wordfile(filename)
answer = read_answers(ansname)
f = open(resultname, "w")
n_problems = len(string_list)
acc = 0
all_ans_acc = 0
correct_ans_acc = 0
# print(string_list)
# print(options_list)
for n in range(n_problems):
	string = string_list[n]
	options = options_list[n]
	n_options = len(options)
	# string = input("string:")
	blank = "XX"
	# n_options = int(input("Number of options:"))
	# options = []


	# for i in range(n_options):
	# 	temp = input(str(i+1) + ":")
	# 	options.append(temp)


	search_string = ""
	a = pos_tag(word_tokenize(string), tagset='universal')
	for i in range(len(a)):
		if a[i][0] == blank:
			target_index = i

	# target_type = input("type:")
	target_type = options_type(options)
	print("type:", target_type)
	a[target_index] = (blank, target_type)

	forward_index = target_index-1
	backward_index = target_index+1

	VERB_index = target_index
	NOUN_index = target_index
	ADJ_index = target_index
	MARKS = [',', '.', ';']
	BeVERB = ['is', 'am', 'are', 'was', 'were']




	if a[target_index][1] == 'ADJ':
		search_string = ADJ_target(forward_index, backward_index)

	elif a[target_index][1] == 'NOUN':
		search_string = NOUN_target(forward_index, backward_index)


	elif a[target_index][1] == 'VERB':
		search_string = VERB_target(forward_index, backward_index)

	elif a[target_index][1] == 'ADV':
		search_string = ADV_target(forward_index, backward_index)
		

	
	options_string = []
	for i in range(n_options):
		options_string.append(search_string.replace(blank, options[i]))
	print(search_string)
	print(options_string)
	print(a)
	search_count = 0
	for option in options_string:
		f.write(option + ",")
	f.write("\n")
	results = []
	for string in options_string:
		result = google_search(string,1)
		# result = random.randint(1, 10)
		results.append(result)
		f.write(str(result)+",")
		# if search_count % 2 == 0:
		# 	t = random.randint(5, 15)
		# elif search_count % 2 == 1:
		# 	t = random.randint(20, 30)
		print(result)
		time.sleep(1)
		search_count += 1
		

	
	

	result_dict = build_result_dict(results)
	ans_pred = max_key(result_dict)
	total = sum(result_dict.values())
	acc_pr = 0
	if total != 0 :
		acc_pr = result_dict[answer[n]] / total
	all_ans_acc += acc_pr	

	if ans_pred == answer[n] and total != 0:
		acc += 1
		correct_ans_acc += acc_pr
		f.write("Y,")
	else:
		f.write("N,")

	f.write("\n")
	print()


print("accuracy: %f  %d/%d" % (acc / n_problems, acc, n_problems))
print("all_ans_acc: %f" % (all_ans_acc / n_problems))
print("correct_ans_acc: %f" % (correct_ans_acc / acc))
f.write("accuracy,%f,%d,%d\n" % (acc / n_problems, acc, n_problems))
f.write("all_ans_acc,{}\n" .format(all_ans_acc / n_problems))
f.write("correct_ans_acc,{}\n" .format(correct_ans_acc / acc))
f.close()

# f = open("result/all.csv", "w")
# f.write("n_problems,acc_problems,all_ans_acc,correct_ans_acc,\n")
# f.write("%d,%d,%f,%f,\n" % (n_problems, acc, all_ans_acc, correct_ans_acc))
# f.close()